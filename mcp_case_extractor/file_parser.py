"""
Document parser for extracting case information from legal documents.

This module handles parsing of various document formats (PDF, DOCX, TXT, HTML)
and extracts structured case information using patterns and LLM.
"""

import asyncio
import logging
import re
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple, Union
from datetime import datetime
import json

# Document processing libraries
try:
    import PyPDF2
    from PyPDF2 import PdfReader
except ImportError:
    PyPDF2 = None
    
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None
    
try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None
    
try:
    import pytesseract
    from PIL import Image
except ImportError:
    pytesseract = None
    Image = None

from openai import AsyncOpenAI

from .models import (
    ExtractedCaseInfo,
    Party,
    PartyType,
    CourtInfo,
    LegalIssue,
    ReliefSought,
    DocumentReference,
    CaseType,
    CaseStage,
    DocumentType
)
from .patterns import LegalPatterns


logger = logging.getLogger(__name__)


class FileParser:
    """Parser for extracting case information from documents."""
    
    SUPPORTED_FORMATS = {'.pdf', '.docx', '.doc', '.txt', '.html', '.htm'}
    
    def __init__(self, openai_api_key: Optional[str] = None, config: Optional[Dict] = None):
        """Initialize the file parser."""
        self.config = config or {}
        self.client = AsyncOpenAI(api_key=openai_api_key) if openai_api_key else None
        self.use_llm = self.client is not None
        self.max_file_size = self.config.get('max_file_size', 10 * 1024 * 1024)  # 10MB default
        self.enable_ocr = self.config.get('enable_ocr', False) and pytesseract is not None
        
    async def parse_document(self, file_path: Union[str, Path]) -> ExtractedCaseInfo:
        """
        Parse a document and extract case information.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            ExtractedCaseInfo object with extracted information
        """
        file_path = Path(file_path)
        
        # Validate file
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if file_path.stat().st_size > self.max_file_size:
            raise ValueError(f"File size exceeds maximum of {self.max_file_size} bytes")
        
        if file_path.suffix.lower() not in self.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported file format: {file_path.suffix}")
        
        # Extract text from document
        text = await self._extract_text(file_path)
        
        if not text or len(text.strip()) < 50:
            raise ValueError("Insufficient text extracted from document")
        
        # Detect document type
        doc_type = LegalPatterns.detect_document_type(text)
        
        # Extract information using patterns
        case_info = self._extract_with_patterns(text)
        
        # Set document metadata
        case_info.extraction_source = "document"
        case_info.document_type = self._map_document_type(doc_type)
        
        # Enhance with LLM if available
        if self.use_llm:
            case_info = await self._enhance_with_llm(text, case_info)
        
        # Calculate confidence score
        case_info.confidence_score = self._calculate_confidence(case_info, text)
        
        return case_info
    
    async def parse_batch(self, file_paths: List[Union[str, Path]]) -> List[ExtractedCaseInfo]:
        """
        Parse multiple documents in batch.
        
        Args:
            file_paths: List of document file paths
            
        Returns:
            List of ExtractedCaseInfo objects
        """
        tasks = [self.parse_document(file_path) for file_path in file_paths]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        extracted_cases = []
        errors = []
        
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                errors.append(f"Error parsing {file_paths[i]}: {str(result)}")
                logger.error(f"Failed to parse {file_paths[i]}: {result}")
            else:
                extracted_cases.append(result)
        
        if errors:
            logger.warning(f"Batch parsing completed with {len(errors)} errors")
        
        return extracted_cases
    
    async def _extract_text(self, file_path: Path) -> str:
        """Extract text from document based on file type."""
        suffix = file_path.suffix.lower()
        
        if suffix == '.txt':
            return self._extract_from_txt(file_path)
        elif suffix == '.pdf':
            return await self._extract_from_pdf(file_path)
        elif suffix in ['.docx', '.doc']:
            return self._extract_from_docx(file_path)
        elif suffix in ['.html', '.htm']:
            return self._extract_from_html(file_path)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")
    
    def _extract_from_txt(self, file_path: Path) -> str:
        """Extract text from TXT file."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    async def _extract_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        if PyPDF2 is None:
            raise ImportError("PyPDF2 is required for PDF parsing")
        
        text = ""
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PdfReader(f)
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num}: {e}")
                        
                        # Try OCR if enabled and text extraction failed
                        if self.enable_ocr:
                            # OCR implementation would go here
                            pass
        except Exception as e:
            logger.error(f"Failed to read PDF file: {e}")
            raise
        
        return text
    
    def _extract_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        if DocxDocument is None:
            raise ImportError("python-docx is required for DOCX parsing")
        
        try:
            doc = DocxDocument(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += "\n" + cell.text
            
            return text
        except Exception as e:
            logger.error(f"Failed to read DOCX file: {e}")
            raise
    
    def _extract_from_html(self, file_path: Path) -> str:
        """Extract text from HTML file."""
        if BeautifulSoup is None:
            raise ImportError("beautifulsoup4 is required for HTML parsing")
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                soup = BeautifulSoup(f.read(), 'html.parser')
                
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()
                
                # Get text
                text = soup.get_text()
                
                # Clean up whitespace
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text = '\n'.join(chunk for chunk in chunks if chunk)
                
                return text
        except Exception as e:
            logger.error(f"Failed to read HTML file: {e}")
            raise
    
    def _extract_with_patterns(self, text: str) -> ExtractedCaseInfo:
        """Extract information using regex patterns."""
        case_info = ExtractedCaseInfo(extraction_source="document")
        
        # Extract case number
        case_number = LegalPatterns.extract_case_number(text)
        if case_number:
            case_info.case_number = case_number
        
        # Extract dates
        dates = LegalPatterns.extract_dates(text)
        if dates:
            # Use the first date as filing date (could be improved with context)
            case_info.filing_date = dates[0][1]
        
        # Extract parties
        parties_dict = LegalPatterns.extract_parties(text)
        for plaintiff_name in parties_dict.get('plaintiffs', []):
            party = Party(
                name=plaintiff_name,
                party_type=PartyType.PLAINTIFF
            )
            case_info.parties.append(party)
        
        for defendant_name in parties_dict.get('defendants', []):
            party = Party(
                name=defendant_name,
                party_type=PartyType.DEFENDANT
            )
            case_info.parties.append(party)
        
        # Extract attorneys
        attorneys = LegalPatterns.extract_attorneys(text)
        # Assign attorneys to parties (simplified - could be improved)
        if attorneys and case_info.parties:
            for party in case_info.parties[:len(attorneys)]:
                party.attorneys = [attorneys[0]] if attorneys else []
                attorneys = attorneys[1:]
        
        # Extract court information
        court_dict = LegalPatterns.extract_court_info(text)
        if court_dict['court']:
            case_info.court_info = CourtInfo(
                name=court_dict['court'],
                jurisdiction=court_dict['jurisdiction'] or 'unknown',
                judge=court_dict['judge']
            )
        
        # Extract citations
        citations = LegalPatterns.extract_citations(text)
        for case_cite in citations.get('cases', []):
            ref = DocumentReference(
                reference_type='case',
                citation=case_cite
            )
            case_info.document_references.append(ref)
        
        for statute_cite in citations.get('statutes', []):
            ref = DocumentReference(
                reference_type='statute',
                citation=statute_cite
            )
            case_info.document_references.append(ref)
        
        # Extract monetary amounts
        amounts = LegalPatterns.extract_monetary_amounts(text)
        if amounts:
            case_info.relief_sought = ReliefSought(
                monetary_damages=max(amounts)  # Use largest amount
            )
        
        # Try to extract case title from parties
        if case_info.parties and len(case_info.parties) >= 2:
            plaintiff_name = case_info.parties[0].name
            defendant_name = case_info.parties[1].name
            case_info.case_title = f"{plaintiff_name} v. {defendant_name}"
        
        return case_info
    
    async def _enhance_with_llm(self, text: str, case_info: ExtractedCaseInfo) -> ExtractedCaseInfo:
        """Enhance extracted information using LLM."""
        if not self.client:
            return case_info
        
        # Truncate text if too long
        max_text_length = 8000
        if len(text) > max_text_length:
            text = text[:max_text_length] + "..."
        
        # Create extraction prompt
        system_prompt = """You are a legal document analyzer. Extract and enhance case information from the provided text.
        Focus on information that wasn't captured by pattern matching.
        Return a JSON object with any additional information found."""
        
        user_prompt = f"""Document text:
        {text}
        
        Already extracted:
        - Case number: {case_info.case_number}
        - Parties: {len(case_info.parties)} found
        - Court: {case_info.court_info.name if case_info.court_info else 'Not found'}
        
        Please extract:
        1. Case type (civil/criminal/family/etc)
        2. Current case stage
        3. Legal issues and claims
        4. Key facts summary
        5. Any disputed facts
        6. Additional parties or attorneys
        7. Relief sought details
        
        Return as JSON with these fields:
        {{
            "case_type": "...",
            "case_stage": "...",
            "legal_issues": [
                {{"issue": "...", "category": "...", "is_primary": true/false}}
            ],
            "fact_summary": "...",
            "disputed_facts": ["..."],
            "additional_parties": [{{"name": "...", "type": "...", "attorneys": []}}],
            "relief_details": "..."
        }}"""
        
        try:
            response = await self.client.chat.completions.create(
                model="gpt-4-turbo-preview",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.1,
                response_format={"type": "json_object"}
            )
            
            enhanced_data = json.loads(response.choices[0].message.content)
            
            # Update case info with enhanced data
            if enhanced_data.get('case_type'):
                case_type_str = enhanced_data['case_type'].lower()
                for case_type in CaseType:
                    if case_type.value == case_type_str:
                        case_info.case_type = case_type
                        break
            
            if enhanced_data.get('case_stage'):
                stage_str = enhanced_data['case_stage'].lower()
                for stage in CaseStage:
                    if stage.value in stage_str:
                        case_info.case_stage = stage
                        break
            
            if enhanced_data.get('legal_issues'):
                for issue_data in enhanced_data['legal_issues']:
                    issue = LegalIssue(
                        issue=issue_data.get('issue', ''),
                        category=issue_data.get('category', 'general'),
                        is_primary=issue_data.get('is_primary', False)
                    )
                    case_info.legal_issues.append(issue)
            
            if enhanced_data.get('fact_summary'):
                case_info.fact_summary = enhanced_data['fact_summary']
            
            if enhanced_data.get('disputed_facts'):
                case_info.disputed_facts = enhanced_data['disputed_facts']
            
            if enhanced_data.get('relief_details'):
                if not case_info.relief_sought:
                    case_info.relief_sought = ReliefSought()
                case_info.relief_sought.other_relief = [enhanced_data['relief_details']]
            
        except Exception as e:
            logger.error(f"LLM enhancement failed: {e}")
        
        return case_info
    
    def _map_document_type(self, doc_type_str: Optional[str]) -> Optional[DocumentType]:
        """Map string document type to DocumentType enum."""
        if not doc_type_str:
            return None
        
        mapping = {
            'complaint': DocumentType.COMPLAINT,
            'answer': DocumentType.ANSWER,
            'motion': DocumentType.MOTION,
            'brief': DocumentType.BRIEF,
            'order': DocumentType.ORDER,
            'judgment': DocumentType.JUDGMENT,
            'notice': DocumentType.NOTICE,
            'discovery': DocumentType.DISCOVERY,
            'pleading': DocumentType.PLEADING,
        }
        
        return mapping.get(doc_type_str, DocumentType.OTHER)
    
    def _calculate_confidence(self, case_info: ExtractedCaseInfo, text: str) -> float:
        """Calculate confidence score for extraction."""
        score = 0.0
        max_score = 0.0
        
        # Score based on field completeness
        field_weights = {
            'case_number': 0.15,
            'case_title': 0.15,
            'parties': 0.20,
            'court_info': 0.15,
            'legal_issues': 0.15,
            'filing_date': 0.10,
            'document_references': 0.10,
        }
        
        for field, weight in field_weights.items():
            max_score += weight
            value = getattr(case_info, field, None)
            
            if value:
                if isinstance(value, list):
                    if len(value) > 0:
                        score += weight
                elif isinstance(value, str):
                    if len(value.strip()) > 0:
                        score += weight
                else:
                    score += weight
        
        # Adjust based on text length (longer documents typically have more info)
        if len(text) > 5000:
            score += 0.1
        
        # Normalize to 0-1 range
        confidence = min(1.0, score / max_score if max_score > 0 else 0)
        
        return confidence